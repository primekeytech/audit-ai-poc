# tests/create_mocks.py
# Run with: python tests/create_mocks.py

import os
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment

def create_mock_policy():
    doc = Document()
    doc.add_heading("First National Bank — IT Security Policy", 0)
    doc.add_heading("1. Access Control Policy", level=1)
    doc.add_paragraph(
        "All user accounts must be reviewed quarterly by IT management. "
        "Privileged access requires dual approval and is reviewed monthly. "
        "Terminated employees must have access revoked within 24 hours. "
        "MFA is required for all remote access and privileged accounts."
    )
    doc.add_heading("2. Change Management Policy", level=1)
    doc.add_paragraph(
        "All system changes must be submitted via the change request system "
        "at least 5 business days in advance. Emergency changes require CISO "
        "approval and must be documented within 24 hours. Changes are tested "
        "in staging before production. A rollback plan is required for all Tier 1 changes."
    )
    doc.add_heading("3. Patch Management", level=1)
    doc.add_paragraph(
        "Critical patches must be applied within 30 days of release. "
        "High-severity patches within 60 days. Systems are scanned monthly. "
        "Exceptions require documented risk acceptance signed by the CISO."
    )
    doc.save("tests/mock_data/bank_it_policy.docx")
    print("✅ Created: bank_it_policy.docx")

def create_mock_questionnaire():
    doc = Document()
    doc.add_heading("IT Audit Questionnaire — Completed by Bank IT Team", 0)
    qa_pairs = [
        (
            "Q1: How often are user access reviews performed?",
            "We perform user access reviews quarterly. Last review was March 31, 2026. "
            "Results are documented and exceptions tracked to remediation."
        ),
        (
            "Q2: Is MFA enforced for privileged accounts?",
            "Yes, MFA is enforced for all admin accounts via Microsoft Authenticator. "
            "VPN access also requires MFA. We have 100% coverage."
        ),
        (
            "Q3: Describe your change management process.",
            "We use ServiceNow for all change requests. Changes go through CAB review "
            "every Tuesday. Emergency changes require CISO sign-off. We had 3 emergency "
            "changes last quarter, all properly documented."
        ),
        (
            "Q4: How are terminated employee accounts handled?",
            "HR notifies IT via automated workflow. Accounts disabled within 2 hours. "
            "Last quarter we had 2 exceptions where accounts were disabled on day 2 "
            "due to weekend terminations."
        ),
        (
            "Q5: What is your patch management timeline for critical patches?",
            "Critical patches applied within 30 days. Currently behind — 4 servers have "
            "a critical patch that is 45 days old due to maintenance window conflict. "
            "Remediation expected by June 15, 2026."
        ),
    ]
    for question, answer in qa_pairs:
        doc.add_paragraph(question, style="Heading 2")
        doc.add_paragraph(answer)
        doc.add_paragraph("")
    doc.save("tests/mock_data/questionnaire_responses.docx")
    print("✅ Created: questionnaire_responses.docx")

def create_mock_workbook():
    wb = Workbook()
    ws = wb.active
    ws.title = "Work Program"

    headers = [
        "Control ID", "Control Domain", "Control Objective",
        "Test Procedure", "Expected Evidence",
        "Evidence Obtained", "AI Analysis", "Score (1-4)"
    ]
    ws.append(headers)

    controls = [
        ["AC-01", "Access Control",
         "User access reviews performed quarterly",
         "1. Obtain access review schedule\n2. Review last 4 quarters\n3. Verify exceptions tracked",
         "Quarterly review reports, sign-off docs, exception log",
         "", "", ""],
        ["AC-02", "Access Control",
         "Privileged accounts reviewed monthly and MFA enforced",
         "1. Obtain privileged account list\n2. Verify MFA enrollment\n3. Review monthly evidence",
         "Privileged account list, MFA screenshots, monthly review logs",
         "", "", ""],
        ["AC-03", "Access Control",
         "Terminated employee accounts disabled within 24 hours",
         "1. Obtain termination list\n2. Compare to disable timestamps\n3. Document exceptions",
         "HR termination report, AD disable logs, exception docs",
         "", "", ""],
        ["CM-01", "Change Management",
         "All changes approved through formal change management process",
         "1. Obtain change log\n2. Verify CAB approval for sample of 25\n3. Test emergency process",
         "Change requests, CAB minutes, CISO approvals",
         "", "", ""],
        ["PM-01", "Patch Management",
         "Critical patches applied within 30 days of release",
         "1. Obtain vulnerability scan results\n2. Identify critical patches\n3. Verify remediation dates",
         "Vulnerability scan reports, patch records, exception docs",
         "", "", ""],
        ["IR-01", "Incident Response",
         "Security incidents reported and escalated per policy",
         "1. Obtain incident log\n2. Verify reporting timelines\n3. Review post-incident reviews",
         "Incident tickets, escalation logs, post-incident reports",
         "", "", ""],
    ]

    for row in controls:
        ws.append(row)

    # Style header row
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="0D1B2A")
    for col in range(1, 9):
        cell = ws.cell(row=1, column=col)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(wrap_text=True, vertical="top")

    # Highlight F, G, H (AI columns)
    ai_fill = PatternFill("solid", fgColor="1A2E44")
    for row in range(2, ws.max_row + 1):
        for col in [6, 7, 8]:
            ws.cell(row=row, column=col).fill = ai_fill

    # Column widths
    for col, width in zip("ABCDEFGH", [10, 18, 35, 40, 35, 35, 40, 12]):
        ws.column_dimensions[col].width = width

    wb.save("tests/mock_data/work_program_template.xlsx")
    print("✅ Created: work_program_template.xlsx")

if __name__ == "__main__":
    os.makedirs("tests/mock_data", exist_ok=True)
    print("\n🔧 Creating mock test files...\n")
    create_mock_policy()
    create_mock_questionnaire()
    create_mock_workbook()
    print("\n✅ Done! Files in tests/mock_data/")
