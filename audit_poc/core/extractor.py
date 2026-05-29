# ============================================
# AUDIT POC - Document Extraction Engine
# ============================================
# This file handles reading and extracting text
# from all document types:
# - PDF files (using PyMuPDF)
# - Word documents (using python-docx)
# - Excel files (using openpyxl)
# ============================================

import os
from pathlib import Path

# PDF extraction library
import fitz  # PyMuPDF

# Word document library
from docx import Document as DocxDocument

# Excel library
import openpyxl

# ============================================
# MAIN EXTRACTION FUNCTION
# ============================================

def extract_all_documents(uploaded_files):
    """
    Main function - extracts text from ALL uploaded files.
    
    Args:
        uploaded_files: List of Streamlit uploaded file objects
        
    Returns:
        dict: {filename: extracted_text} for each file
    
    Example:
        {
            "SecurityPolicy.pdf": "This policy covers...",
            "ChangeControl.docx": "Change control process...",
        }
    """
    
    # Dictionary to store all extracted text
    extracted = {}
    
    # Loop through each uploaded file
    for file in uploaded_files:
        
        # Get the file name and extension
        filename = file.name
        extension = Path(filename).suffix.lower()  # .pdf, .docx, .xlsx
        
        print(f"Extracting: {filename} ({extension})")
        
        try:
            # Route to correct extractor based on file type
            if extension == ".pdf":
                # Extract text from PDF
                text = extract_pdf(file)
                
            elif extension == ".docx":
                # Extract text from Word document
                text = extract_docx(file)
                
            elif extension == ".xlsx":
                # Extract text from Excel file
                text = extract_xlsx(file)
                
            else:
                # Unknown file type - skip it
                print(f"Skipping unsupported file type: {extension}")
                text = ""
            
            # Store extracted text with filename as key
            extracted[filename] = text
            print(f"Successfully extracted {len(text)} characters from {filename}")
            
        except Exception as e:
            # If extraction fails - log error and continue
            print(f"Error extracting {filename}: {str(e)}")
            extracted[filename] = f"ERROR: Could not extract text - {str(e)}"
    
    return extracted

# ============================================
# PDF EXTRACTION
# ============================================

def extract_pdf(file):
    """
    Extracts all text from a PDF file.
    Uses PyMuPDF (fitz) for fast reliable extraction.
    
    Args:
        file: Streamlit uploaded file object
        
    Returns:
        str: All text extracted from the PDF
    """
    
    # Read file bytes from Streamlit uploader
    file_bytes = file.read()
    
    # Open PDF from bytes using PyMuPDF
    pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
    
    # Store all text from all pages
    all_text = []
    
    # Loop through every page in the PDF
    for page_number in range(len(pdf_document)):
        
        # Get the page
        page = pdf_document[page_number]
        
        # Extract text from this page
        page_text = page.get_text()
        
        # Add page header so we know which page the text came from
        all_text.append(f"\n--- PAGE {page_number + 1} ---\n")
        all_text.append(page_text)
    
    # Close the PDF document
    pdf_document.close()
    
    # Join all text into one string
    return "\n".join(all_text)

# ============================================
# WORD DOCUMENT EXTRACTION
# ============================================

def extract_docx(file):
    """
    Extracts all text from a Word document (.docx).
    Uses python-docx library.
    
    Args:
        file: Streamlit uploaded file object
        
    Returns:
        str: All text extracted from the Word document
    """
    
    # Open the Word document
    doc = DocxDocument(file)
    
    # Store all paragraphs
    all_text = []
    
    # Loop through every paragraph in the document
    for paragraph in doc.paragraphs:
        
        # Only add non-empty paragraphs
        if paragraph.text.strip():
            all_text.append(paragraph.text)
    
    # Also extract text from tables in the document
    for table in doc.tables:
        
        # Loop through each row in the table
        for row in table.rows:
            
            # Get text from each cell in the row
            row_text = " | ".join(
                cell.text.strip() 
                for cell in row.cells 
                if cell.text.strip()
            )
            
            # Add row text if not empty
            if row_text:
                all_text.append(row_text)
    
    # Join all text with newlines
    return "\n".join(all_text)

# ============================================
# EXCEL EXTRACTION
# ============================================

def extract_xlsx(file):
    """
    Extracts all text/values from an Excel file (.xlsx).
    Uses openpyxl library.
    
    IMPORTANT: Uses data_only=True here because we are
    READING values from questionnaire/evidence files.
    
    For the WORK PROGRAM template - we use a different
    function (in workbook_engine.py) with data_only=False
    to preserve formulas!
    
    Args:
        file: Streamlit uploaded file object
        
    Returns:
        str: All cell values extracted from Excel
    """
    
    # Load workbook - data_only=True means read cell VALUES not formulas
    # This is correct for questionnaire/evidence files
    workbook = openpyxl.load_workbook(file, data_only=True)
    
    # Store all text from all sheets
    all_text = []
    
    # Loop through every sheet in the workbook
    for sheet_name in workbook.sheetnames:
        
        # Get the sheet
        sheet = workbook[sheet_name]
        
        # Add sheet name as header
        all_text.append(f"\n--- SHEET: {sheet_name} ---\n")
        
        # Loop through every row in the sheet
        for row in sheet.iter_rows():
            
            # Get text from each cell in the row
            row_values = []
            for cell in row:
                
                # Only include non-empty cells
                if cell.value is not None:
                    row_values.append(str(cell.value))
            
            # Add row if it has any values
            if row_values:
                all_text.append(" | ".join(row_values))
    
    # Close the workbook
    workbook.close()
    
    # Join all text with newlines
    return "\n".join(all_text)

# ============================================
# HELPER FUNCTION - COMBINE ALL TEXT
# ============================================

def combine_extracted_text(extracted_dict):
    """
    Combines all extracted text from multiple documents
    into one single context string for the AI.
    
    Args:
        extracted_dict: {filename: text} dictionary
        
    Returns:
        str: Single combined context string
    """
    
    combined = []
    
    # Add each document's text with a clear separator
    for filename, text in extracted_dict.items():
        combined.append(f"\n{'='*50}")
        combined.append(f"DOCUMENT: {filename}")
        combined.append(f"{'='*50}\n")
        combined.append(text)
    
    return "\n".join(combined)
